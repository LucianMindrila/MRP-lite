from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ─────────────────────────────────────────────────────
DARK  = RGBColor(0x1a, 0x1f, 0x2e)
BLUE  = RGBColor(0x0d, 0x6e, 0xfd)
GREEN = RGBColor(0x19, 0x87, 0x54)
AMBER = RGBColor(0xfd, 0x7e, 0x14)
GREY  = RGBColor(0x6c, 0x75, 0x7d)
WHITE = RGBColor(0xff, 0xff, 0xff)

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    hex_col = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:fill'),  hex_col)
    tcPr.append(shd)

def heading(text, level=1, colour=DARK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = colour
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False, colour=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    if bold or colour:
        for run in p.runs:
            if bold:   run.bold = True
            if colour: run.font.color.rgb = colour
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A1F2E')
    pBdr.append(bot)
    pPr.append(pBdr)

def section_box(label, colour: RGBColor):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, colour)
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(11)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading('MRP Lite — Operator Workflow Plan', 0)
for run in title.runs:
    run.font.color.rgb = DARK
title.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph('Goods In & Goods Out — Tablet, Barcode Scanner & Label Printer')
sub.runs[0].font.color.rgb = GREY
sub.runs[0].italic = True
sub.paragraph_format.space_after = Pt(2)

date_p = doc.add_paragraph('Prepared: 08 May 2026')
date_p.runs[0].font.size = Pt(9)
date_p.runs[0].font.color.rgb = GREY

divider()

# ═══════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ═══════════════════════════════════════════════════════════════
heading('1. What We Are Building', 1)
body('Three new capabilities will be added on top of the existing MRP Lite admin system:')
bullet('A tablet UI — a separate, touch-optimised interface the operator sees when they log in. No sidebar, large buttons, simple focused screens. The admin continues to use the existing full desktop interface.')
bullet('Barcode integration — labels printed from the browser, scanning done via any Bluetooth or USB scanner. Scanners work as keyboard input — they simply "type" the barcode value into whatever input field is currently focused. No special drivers or apps required.')
bullet('Two operator workflows — Goods In and Goods Out.')

# ═══════════════════════════════════════════════════════════════
# SECTION 2 — HARDWARE
# ═══════════════════════════════════════════════════════════════
heading('2. Hardware Setup (One-Time)', 1)
body('All three devices connect over WiFi. No special software installation is needed — everything runs through the browser on the tablet pointing at the MRP Lite server on your network.')
doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Device', 'Recommendation', 'Why']
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, DARK)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

rows_data = [
    ('Tablet',          'iPad (any modern) or Android 10"+',           'Browser-based, no app install needed'),
    ('Barcode scanner', 'Bluetooth handheld — Zebra CS60 or Honeywell Voyager', 'Connects to tablet as keyboard input'),
    ('Label printer',   'Brother QL-820NWB or Zebra ZD220',            'Prints direct from browser via WiFi'),
]
for i, (dev, rec, why) in enumerate(rows_data, start=1):
    row = tbl.rows[i]
    row.cells[0].paragraphs[0].add_run(dev).bold = True
    row.cells[1].paragraphs[0].add_run(rec)
    row.cells[2].paragraphs[0].add_run(why)
    if i % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, RGBColor(0xf4, 0xf6, 0xf9))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# SECTION 3 — FLOW 1: GOODS IN
# ═══════════════════════════════════════════════════════════════
heading('3. Flow 1 — Goods In', 1, DARK)
divider()

heading('3.1  Admin Side (Desktop — Done in Advance)', 2, BLUE)
numbered('Admin raises Purchase Orders in MRP Lite as normal — supplier, materials, quantities, expected delivery date.')
numbered('PO status is set to Draft, then updated to Sent once issued to the supplier.')
numbered('That is all the admin needs to do. The PO is now visible to the operator on their tablet.')

heading('3.2  Operator Side (Tablet — Day of Delivery)', 2, GREEN)

body('Step 1 — Open the tablet', bold=True)
body('The operator logs in. Because their account role is set to Operator, they land directly on the Operator Dashboard — a full-screen touch interface with two large tiles:')

code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Inches(0.5)
run = code_p.add_run(
    'GOODS IN               GOODS OUT\n'
    '3 expected this week   2 due today'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = DARK

body('Step 2 — Tap Goods In', bold=True)
body('The operator sees a list of open Purchase Orders grouped by expected delivery date. Today\'s deliveries are highlighted at the top.')

body('Step 3 — Identify the delivery', bold=True)
body('Two options are available:')
bullet('Tap the PO from the list if they know which supplier it is.')
bullet('Scan the barcode on the supplier\'s paperwork — if it contains a PO reference the system recognises, it auto-selects the correct PO.')

body('Step 4 — Receive the delivery', bold=True)
body('The PO detail screen shows each line with the expected quantity and a quantity-received input field. The operator enters the quantity received for each line. If a quantity differs from what was expected, the line is highlighted in amber. Short deliveries remain open; surplus is flagged for admin review.')

body('Step 5 — Assign a warehouse location', bold=True)
body('After entering a received quantity, a location field appears. The operator enters or scans the location code for the shelf or bay where the stock is being put away (for example BAY-A3). Location barcodes can be printed and stuck on shelving in the warehouse.')

body('Step 6 — Print shelf labels', bold=True)
body('One tap per line sends a label to the label printer. Each label contains:')
bullet('Material code and description')
bullet('Quantity received')
bullet('Warehouse location')
bullet('Date received')
bullet('Barcode of the material code for future scanning')

body('Step 7 — Confirm receipt', bold=True)
body('The operator taps the green Confirm Goods In button. This triggers the following automatically:')
bullet('Stock quantity updated on each material received')
bullet('Stock Movement records created (type: Goods In) with operator name and date')
bullet('Purchase Order lines updated with qty_received values')
bullet('If all lines are fully received, PO status changes to Received')
bullet('If partial delivery, PO status changes to Partial and outstanding lines remain open')
bullet('Admin can see the updated stock and PO status immediately on their desktop')

# ═══════════════════════════════════════════════════════════════
# SECTION 4 — FLOW 2: GOODS OUT
# ═══════════════════════════════════════════════════════════════
heading('4. Flow 2 — Goods Out', 1, DARK)
divider()

heading('4.1  Admin Side (Desktop — Done in Advance)', 2, BLUE)
numbered('Admin creates Delivery Notes as normal — single order dispatch or multi-order dispatch.')
numbered('Each DN has a planned dispatch date. Status is set to Pending.')
numbered('The DN is now visible to the operator grouped by dispatch date.')
numbered('Admin does nothing further until they see the DN status update to Dispatched on their screen.')

heading('4.2  Operator Side (Tablet — Day of Dispatch)', 2, GREEN)

body('Step 1 — Tap Goods Out', bold=True)
body('From the operator dashboard the operator sees Delivery Notes grouped by date:')

code_p2 = doc.add_paragraph()
code_p2.paragraph_format.left_indent = Inches(0.5)
run2 = code_p2.add_run(
    'TODAY\n'
    '  DN-00041   Safety Knife Co    3 lines\n'
    '  DN-00042   NuCo Tools         1 line\n\n'
    'TOMORROW\n'
    '  DN-00043   Trend              2 lines'
)
run2.font.name = 'Courier New'
run2.font.size = Pt(9)
run2.font.color.rgb = DARK

body('Step 2 — Open a Delivery Note', bold=True)
body('The operator taps a DN to open the pick list. They see every item that needs to come off the shelf, including the warehouse location so they know exactly where to go:')

code_p3 = doc.add_paragraph()
code_p3.paragraph_format.left_indent = Inches(0.5)
run3 = code_p3.add_run(
    'DN-00041  |  Safety Knife Company\n'
    '-----------------------------------------------\n'
    '[ ] BIG FISH 9MM N/H     Qty: 500    LOC: BAY-A3\n'
    '[ ] BC-REAKTA-HD 001     Qty: 200    LOC: BAY-B1\n'
    '[ ] 13154002             Qty: 100    LOC: BAY-C2\n'
    '-----------------------------------------------\n'
    '[Print Delivery Note]  [Mark All Dispatched]'
)
run3.font.name = 'Courier New'
run3.font.size = Pt(9)
run3.font.color.rgb = DARK

body('Step 3 — Scan to confirm (optional but recommended)', bold=True)
body('As the operator picks each item off the shelf, they scan its barcode label. The system matches the scan to the correct line and ticks the checkbox automatically. This prevents picking the wrong product and gives a clear record of what was actually loaded.')

body('Step 4 — Print the Delivery Note', bold=True)
body('The operator taps Print Delivery Note. The A4 delivery note is sent to the printer or opens in the browser print dialog. They print it, fold it, and place it with the goods on the lorry.')

body('Step 5 — Mark as Dispatched', bold=True)
body('Once everything is loaded and all checkboxes are ticked, the operator taps Mark All Dispatched. This triggers:')
bullet('Dispatched quantities updated on each order item')
bullet('Stock Movement records created (type: Goods Out) with operator name and date')
bullet('DN status updated to Dispatched')
bullet('If all items on the parent order are now fully dispatched, order status changes to Dispatched')
bullet('Admin sees the status change immediately on their desktop')

# ═══════════════════════════════════════════════════════════════
# SECTION 5 — TECHNICAL BUILD PLAN
# ═══════════════════════════════════════════════════════════════
heading('5. Technical Build Plan', 1, DARK)
divider()
body('The build is split into four phases. Phases A and B deliver Goods In. Phase C delivers Goods Out. Phase D polishes documents before they go to customers.')

heading('Phase A — Tablet UI Foundation', 2, BLUE)
numbered('Add CSS breakpoints for tablet screens — larger touch targets, no sidebar, full-width cards.')
numbered('Create a separate operator template folder with a minimal header (logo, role name, logout only).')
numbered('Modify the login redirect: Operator role goes to /operator/dashboard. Admin and Manager go to the existing full dashboard.')

heading('Phase B — Goods In Flow', 2, GREEN)
numbered('Operator dashboard with Goods In and Goods Out tiles showing live counts.')
numbered('Goods In PO list — open POs sorted by expected date, highlighted if due today.')
numbered('PO receive screen — quantity input per line, location field, scanner-ready inputs.')
numbered('Confirm receipt route — updates stock quantities, creates Stock Movement records, updates PO status.')
numbered('Shelf label template — print-optimised, includes barcode image generated in the browser.')

heading('Phase C — Goods Out Flow', 2, AMBER)
numbered('Goods Out DN list — pending DNs grouped by planned dispatch date.')
numbered('DN pick list screen — items with locations, tick-off checkboxes, scan-to-confirm.')
numbered('Mark dispatched route — updates order items, creates Stock Movement records, updates DN and order status.')
numbered('Wire the print button to the existing delivery note print template.')

heading('Phase D — Document Polish', 2, GREY)
numbered('Move company name, address, and logo to a central config so all documents update from one place.')
numbered('Apply your existing delivery note and invoice formatting — logo, fonts, layout, address block.')
numbered('Apply consistent formatting across all print templates: Delivery Note, Invoice, Work Order.')
numbered('Real-time badge counts on operator dashboard refreshing every 60 seconds.')

# ═══════════════════════════════════════════════════════════════
# SECTION 6 — DECISIONS NEEDED
# ═══════════════════════════════════════════════════════════════
heading('6. Decisions Needed Before We Build', 1, DARK)
divider()
body('Please confirm the following before development starts on the operator flows:')

bullet('Label printer model', bold_prefix='Label printer:')
body('     Brother QL-820NWB and Zebra ZD220 are both compatible. Label size and setup differs slightly between them. Confirm which you will purchase so labels are designed to the right dimensions.')

bullet('WiFi coverage in the warehouse', bold_prefix='WiFi:')
body('     The tablet needs a reliable WiFi signal at the goods-in door and in the dispatch area. If there are dead spots these will need to be addressed before go-live.')

bullet('Individual or shared operator logins', bold_prefix='Operator accounts:')
body('     Individual accounts are strongly recommended. Every Stock Movement record captures who received or dispatched the goods. This is important for traceability and resolving discrepancies. Shared logins lose this audit trail.')

bullet('Location codes', bold_prefix='Warehouse locations:')
body('     We need a simple location coding system (e.g. BAY-A1, SHELF-B3) agreed in advance. Once agreed, location barcodes can be printed and stuck on the shelving. The system will store and display locations per material.')

# ═══════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════
divider()
footer = doc.add_paragraph('MRP Lite — DT Solutions Ltd   |   Prepared by Axiom   |   08 May 2026')
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = GREY
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = r'C:\Users\conta\OneDrive\Desktop\Operator_Workflow_Plan.docx'
doc.save(out)
print('Saved:', out)
