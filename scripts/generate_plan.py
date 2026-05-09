"""Generate MRP Lite implementation plan as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ───────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1a, 0x3a, 0x5c)
MID_BLUE   = RGBColor(0x2e, 0x6d, 0xa4)
LIGHT_BLUE = RGBColor(0xd6, 0xe8, 0xf7)
GREEN      = RGBColor(0x1a, 0x7a, 0x3c)
LIGHT_GREEN= RGBColor(0xd4, 0xed, 0xda)
AMBER      = RGBColor(0x7a, 0x4f, 0x00)
LIGHT_AMBER= RGBColor(0xff, 0xf3, 0xcd)
GREY_BG    = RGBColor(0xf2, 0xf4, 0xf7)
WHITE      = RGBColor(0xff, 0xff, 0xff)
BLACK      = RGBColor(0x1a, 0x1a, 0x1a)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color='CCCCCC'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ── Header ───────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('MRP Lite — Implementation Plan')
run.bold      = True
run.font.size = Pt(22)
run.font.color.rgb = DARK_BLUE

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_before = Pt(0)
sub_para.paragraph_format.space_after  = Pt(2)
run = sub_para.add_run(f'DT Solutions Ltd   ·   {datetime.date.today().strftime("%d %B %Y")}')
run.font.size      = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer


# ── Phase builder ─────────────────────────────────────────────────────────────
def add_phase(title, subtitle, header_rgb, header_light_rgb, rows):
    """Add a phase heading + table."""
    # Phase title
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(10)
    ph.paragraph_format.space_after  = Pt(4)
    r = ph.add_run(title)
    r.bold            = True
    r.font.size       = Pt(13)
    r.font.color.rgb  = header_rgb

    # Sub-description
    if subtitle:
        sd = doc.add_paragraph()
        sd.paragraph_format.space_before = Pt(0)
        sd.paragraph_format.space_after  = Pt(6)
        sr = sd.add_run(subtitle)
        sr.font.size      = Pt(10)
        sr.italic         = True
        sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Table: #  |  Action  |  Who
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style            = 'Table Grid'
    tbl.alignment        = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit    = True

    # Column widths
    widths = [Cm(1.4), Cm(12.0), Cm(3.0)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # Header row
    hdr = tbl.rows[0]
    for i, text in enumerate(['#', 'Action', 'Who']):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_rgb)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold            = True
        run.font.size       = Pt(10)
        run.font.color.rgb  = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, (ref, action, who, status) in enumerate(rows):
        row   = tbl.rows[idx + 1]
        bg    = header_light_rgb if idx % 2 == 0 else WHITE

        # Status overrides bg colour
        if status == 'done':
            bg = LIGHT_GREEN
        elif status == 'auto':
            bg = LIGHT_AMBER

        for cell in row.cells:
            set_cell_bg(cell, bg)
            set_cell_borders(cell, 'CCCCCC')

        # # col
        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        r0 = c0.paragraphs[0].add_run(ref)
        r0.font.size      = Pt(9)
        r0.bold           = True
        r0.font.color.rgb = header_rgb
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Action col
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        r1 = c1.paragraphs[0].add_run(action)
        r1.font.size      = Pt(10)
        r1.font.color.rgb = BLACK

        # Who col
        c2 = row.cells[2]
        c2.paragraphs[0].clear()
        who_color = GREEN if who == 'Done ✓' else (AMBER if who == 'Automatic' else (MID_BLUE if who == 'Axiom' else BLACK))
        r2 = c2.paragraphs[0].add_run(who)
        r2.font.size      = Pt(10)
        r2.bold           = (who == 'Done')
        r2.font.color.rgb = who_color
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer after table


# ── Phase 1 ──────────────────────────────────────────────────────────────────
add_phase(
    'Phase 1 — Product & BOM Completeness',
    'Establish what you make and exactly what goes into each product.',
    MID_BLUE, LIGHT_BLUE,
    [
        ('1.1', 'Add BOM status indicator to the Products list — red badge (no BOM) / green badge (has BOM)', 'Axiom',  ''),
        ('1.2', 'Work through every product and enter its full BOM — components and quantity per unit',        'Lucian', ''),
        ('1.3', 'Verify BOM cost per product is calculating correctly across all products',                    'Axiom & Lucian',''),
    ]
)

# ── Phase 2 ──────────────────────────────────────────────────────────────────
add_phase(
    'Phase 2 — Materials & Supplier Data',
    'Know what you buy, who you buy it from, what it costs, and when it arrives.',
    RGBColor(0x5a, 0x35, 0x8c), RGBColor(0xec, 0xe6, 0xf7),
    [
        ('2.1', 'Add all suppliers to the Suppliers tab — name, contact details, lead time in days',           'Lucian', ''),
        ('2.2', 'Assign a supplier to every material in the Materials tab',                                    'Lucian', ''),
        ('2.3', 'Enter cost price on every material',                                                          'Lucian', ''),
        ('2.4', 'Set reorder point (minimum stock level) and reorder quantity on each material',               'Lucian', ''),
    ]
)

# ── Phase 3 ──────────────────────────────────────────────────────────────────
add_phase(
    'Phase 3 — Initial Stock Count',
    'Establish an accurate baseline of what is physically on the shelf.',
    RGBColor(0x1a, 0x6e, 0x4a), RGBColor(0xd4, 0xed, 0xda),
    [
        ('3.1', 'Build a Stock Count screen — all materials on one page with quantity input fields',           'Axiom',        ''),
        ('3.2', 'Physical warehouse count — walk every shelf and record actual quantities',                    'Lucian',       ''),
        ('3.3', 'Enter counts into the Stock Count screen',                                                    'Lucian',       ''),
        ('3.4', 'System records the stock adjustment and updates all live quantities',                         'Automatic', 'auto'),
    ]
)

# ── Phase 4 ──────────────────────────────────────────────────────────────────
add_phase(
    'Phase 4 — Purchase Requirements Report',
    'Calculate exactly what needs to be ordered and generate purchase orders automatically.',
    RGBColor(0x8a, 0x3a, 0x1a), RGBColor(0xf9, 0xe8, 0xe0),
    [
        ('4.1', 'Build "What to Buy" report — confirmed orders × BOM minus stock on hand = shortfall by material grouped by supplier', 'Axiom',  ''),
        ('4.2', 'Review shortfall report against business knowledge and adjust if needed',                     'Lucian', ''),
        ('4.3', 'Build one-click conversion of shortfall into a draft Purchase Order per supplier',            'Axiom',  ''),
        ('4.4', 'Review, adjust quantities/prices, and send purchase orders to suppliers',                     'Lucian', ''),
    ]
)

# ── Phase 5 ──────────────────────────────────────────────────────────────────
add_phase(
    'Phase 5 — Ongoing Operations',
    'Day-to-day running once the system is fully populated.',
    RGBColor(0x2a, 0x2a, 0x2a), RGBColor(0xf0, 0xf0, 0xf0),
    [
        ('5.1', 'New orders auto-imported from Outlook via Power Automate + watcher (max 10 min)',             'Done ✓',    'done'),
        ('5.2', 'Stock automatically deducted when a delivery note is created',                                'Done ✓',    'done'),
        ('5.3', 'Goods in recorded via Purchase Order receiving screen',                                       'Done ✓',    'done'),
        ('5.4', 'Build reorder alerts — flag materials that have dropped below their reorder point',           'Axiom',        ''),
    ]
)

# ── Legend ───────────────────────────────────────────────────────────────────
leg = doc.add_paragraph()
leg.paragraph_format.space_before = Pt(6)
leg.add_run('Key:  ').bold = True

def legend_chip(para, color, label):
    r = para.add_run(f'  {label}  ')
    r.font.color.rgb  = color
    r.font.size       = Pt(9)
    r.bold            = True

legend_chip(leg, GREEN,              'Done')
leg.add_run('  ')
legend_chip(leg, AMBER,              'Automatic')
leg.add_run('  ')
legend_chip(leg, MID_BLUE,           'Axiom')
leg.add_run('  ')
legend_chip(leg, BLACK,              'Lucian')

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\conta\OneDrive\Desktop\MRP_Lite_Implementation_Plan.docx'
doc.save(out)
print(f"Saved: {out}")
